import csv
import re
import markdown
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import requests
import json
from docx import Document
import os
import uuid
import mistune
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString
from docx.document import Document as DocumentObject

# Load environment variables
from dotenv import load_dotenv
load_dotenv()


# Load Gemini API key
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

# Google Drive API setup
SERVICE_ACCOUNT_FILE = 'service_account.json'
SCOPES = ['https://www.googleapis.com/auth/drive.file']
GOOGLE_DRIVE_FOLDER_ID = os.getenv('GOOGLE_DRIVE_FOLDER_ID')

credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES
)
drive_service = build('drive', 'v3', credentials=credentials)

# Profile object
with open('profile.json') as f:
    profile = json.load(f)


def markdown_to_text(md_string):
    html = mistune.markdown(md_string)
    if not isinstance(html, str):
        html = str(html)
    return BeautifulSoup(html, "html.parser").get_text()


def add_formatted_run(paragraph, element):
    """
    Recursively processes HTML elements and adds styled runs to the paragraph.
    """
    if isinstance(element, NavigableString):
        paragraph.add_run(str(element))
        return

    if not isinstance(element, Tag):
        return

    run = paragraph.add_run()

    if element.name in ['strong', 'b']:
        run.bold = True
    if element.name in ['em', 'i']:
        run.italic = True
    if element.name == 'a':
        run.underline = True
        # You can extend this to add hyperlinks using python-docx custom XML

    for child in element.contents:
        if isinstance(child, NavigableString):
            run.add_text(str(child))
        elif isinstance(child, Tag):
            add_formatted_run(paragraph, child)


def add_formatted_paragraph(doc, element):
    paragraph = doc.add_paragraph()
    for child in element.contents:
        add_formatted_run(paragraph, child)


def addPara(doc: DocumentObject, space=1, text="", style=None):
    paragraph = doc.add_paragraph(text, style=style)
    # paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_after = Pt(
        space)  # or Pt(0) for no extra space
    paragraph.paragraph_format.space_before = Pt(0)
    return paragraph


def markdown_to_docx(markdown_text, output_path):
    print(f"Converting markdown to docx: {output_path}")

    # html = mistune.markdown(markdown_text)
    # print(html)

    # soup = BeautifulSoup(html, 'html.parser')

    # print("Parsed HTML content")

    print("Creating new Document to f{output_path}")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1

    for line in markdown_text.strip().splitlines():

        if line.strip().startswith("***") and line.strip().endswith("***"):
            paragraph = addPara(doc)
            run = paragraph.add_run(line.strip().replace("***", ""))
            run.bold = True
            # run.font.size = Pt(15)

        elif "***" in line:
            parts = line.split("***")
            if len(parts) >= 3:
                # parts[1] is the bold text, parts[2] is the rest
                paragraph = addPara(doc)
                run_bold = paragraph.add_run(parts[1] + " ")
                run_bold.bold = True
                # run_bold.font.size = Pt(15)
                paragraph.add_run(parts[2].strip())

        elif "**" in line:
            parts = line.split("**")
            if len(parts) >= 3:
                # parts[1] is the bold text, parts[2] is the rest
                paragraph = addPara(doc, 6)
                run_bold = paragraph.add_run(parts[1] + " ")
                run_bold.bold = True
                run_bold.font.size = Pt(18)
                paragraph.add_run(parts[2].strip())

        elif line.strip().startswith("**") and line.strip().endswith("**"):
            paragraph = addPara(doc, 6)
            run = paragraph.add_run(line.strip().replace("**", ""))
            run.bold = True
            run.font.size = Pt(18)

        elif line.strip().startswith("### "):
            paragraph = addPara(doc, 3)
            run = paragraph.add_run(line.replace("### ", "").strip())
            run.bold = True
            run.font.size = Pt(15)

        elif line.strip().startswith("- "):
            addPara(doc, 0, line.strip()[1:], style="List Bullet")
        elif line.strip().startswith("---"):
            addPara(doc, 6)
        elif line.strip().startswith("_") and line.strip().endswith("_"):
            paragraph = addPara(doc)
            run = paragraph.add_run(line.strip().strip("_"))
            run.italic = True

        else:
            doc.add_paragraph(line.strip())

    # for element in soup.contents:
    #     if isinstance(element, NavigableString):
    #         doc.add_paragraph(str(element))
    #         continue

    #     # print(element.name)

    #     if element.name == 'h1':
    #         doc.add_heading(element.get_text(), level=1)
    #     elif element.name == 'h2':
    #         doc.add_heading(element.get_text(), level=2)
    #     elif element.name == 'h3':
    #         doc.add_heading(element.get_text(), level=3)
    #     elif element.name == 'ul':
    #         for li in element.find_all('li', recursive=False):
    #             doc.add_paragraph(li.get_text(), style='List Bullet')
    #     elif element.name == 'ol':
    #         for li in element.find_all('li', recursive=False):
    #             doc.add_paragraph(li.get_text(), style='List Number')
    #     elif element.name == 'blockquote':
    #         p = doc.add_paragraph(f'"{element.get_text()}"')
    #         p.paragraph_format.left_indent = Pt(20)
    #         p.runs[0].italic = True
    #     elif element.name == 'pre':
    #         code_text = element.get_text()
    #         doc.add_paragraph(code_text, style='Intense Quote')
    #     elif element.name == 'p':
    #         add_formatted_paragraph(doc, element)
    #     else:
    #         add_formatted_paragraph(doc, element)

    doc.save(output_path)
    print(f"Saved docx to {output_path}")


def sendToGemini(description, company_name=None) -> str | None:
    format = """
Generate a resume in with the following structure and formatting without quoting:

---
**[Full Name]**

***[Job Title]***

[Email Address] | [Github] | Remote

---

### Summary

A short professional summary (3–5 sentences) outlining years of experience, domain expertise, core technologies, and what sets you apart professionally.

---

### Core Skills and Expertise

***Programming Languages:*** [e.g. Lua, C++, Python, JavaScript, JSON, Java, TypeScript, Go]  
***Frameworks & Libraries:*** [e.g. React, Node.js, Qt, Angular, Vue.js, Express, Flask]  
***Development Tools:*** [e.g. Git, JIRA, Docker, Jenkins, VSCode]  
***Web & Mobile Technologies:*** [e.g. RESTful APIs, GraphQL, HTML5, CSS3, Electron]  
***Software & Systems:*** [e.g. Q-SYS, AWS, Azure, Linux, Agile, CI/CD Pipelines]  
***Community & Engagement:*** [e.g. Blogging, Mentorship, Community Management, Speaking]

---

### Education

***[Degree]***   

[Start Date] – [End Date]
---

### Experience

***[Company Name]***

***[Job Title]***

_[Start Date] – [End Date]_

- Bullet point for key responsibility or achievement.
- Use action verbs to start each line.
- Focus on impact, results, and technologies used.

_Repeat as needed for each role._

---

### Additional Skills

***Leadership & Team Management:*** [Details]  
***Agile Methodologies:*** [Details]  
***Communication & Client Engagement:*** [Details]  
***Problem-Solving & Innovation:*** [Details]

---

### Key Projects
***[Organization]***

- Bullet point for each project.
- Use action verbs to start each line.
- Focus on impact, results, and technologies used.

"""
    prompt = (
        "Write a resume matching the requirements and for the above job description in text with no additional tips, include company's name and job role, using my professional informations:\n\n" + json.dumps(profile) +
        "\n\nto tailor the resume. You are allowed to tweak my skills and experience details to suit "
        "the job description. You are allowed to tweak my skills and experience details to suit the job description.\n\n" + format
    )

    payload = {
        "contents": [{"parts": [{"text": f"{description}\n\n{prompt}"}]}]
    }

    # return GEMINI_API_KEY
    response = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}",
        headers={"Content-Type": "application/json"},
        data=json.dumps(payload)
    )

    result = response.json()
    reply = result.get('candidates', [{}])[0].get('content', {}).get(
        'parts', [{}])[0].get('text', None)

    # html = markdown.markdown(reply)

    if reply is None:
        return None
    # Convert markdown to text
    # reply = markdown_to_text(reply)
    # return reply

    # Write to DOCX

    name = company_name if company_name is not None else f"Resume-{uuid.uuid4()}"
    filename = f"{name}.docx"
    doc_path = f"/tmp/resume.docx"
    # print(f"Saving to {doc_path}")
    drive_file_name = filename  # f"{name}_{uuid.uuid4()}.docx"
    markdown_to_docx(reply, doc_path)
    # return None
    # return doc_path

    # Upload to Google Drive
    print(f"Uploading to Google Drive: {drive_file_name}")
    file_metadata = {
        'name': drive_file_name,
        'parents': [GOOGLE_DRIVE_FOLDER_ID],
        'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    media = MediaFileUpload(
        doc_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = drive_service.files().create(
        body=file_metadata, media_body=media, fields='id').execute()

    file_id = file.get('id')

    print(f"File ID: {file_id}")

    permission = {
        'type': 'anyone',
        'role': 'reader',
    }
    drive_service.permissions().create(fileId=file_id, body=permission).execute()
    file_url = f"https://drive.google.com/file/d/{file_id}/view"

    print(f"File URL: {file_url}")

    return file_url


def main():
    start_id = None

    end_id = None

    id_started = False

    path = 'jobs.json'
    all_jobs = {}

    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            all_jobs = json.load(f)
    else:
        all_jobs = {}

    job_count = len(all_jobs)
    print(f"Loaded {job_count} jobs from {path}")
    current_count = 0
    for index in all_jobs:
        current_count += 1
        print(f"{current_count}/{job_count} - {index}")

        if start_id and index != start_id and not id_started:
            # print(f"Skipping job {uid} until start_id {start_id} is found.")
            continue

        if end_id and index == end_id:
            print(f"Found end_id {end_id}, stopping processing.")
            break

        else:
            if not id_started:
                print(f"Found start_id {start_id}, starting to process jobs.")
                id_started = True

        job = all_jobs[index]
        title = job["title"].strip()
        company = job["company"].strip()
        description = job["description"].strip()
        file_url = job["file_url"] if "file_url" in job else None

        print()
        print(f"Processing job at index {index} for {title} at {company}...")

        if file_url is None:
            job_description = f"{title} at {company}\n\n{description}"
            title = re.sub(r'[^a-zA-Z0-9]', '_', title)
            resume_link = sendToGemini(job_description, f"{company} - {title}")

            if resume_link is None:
                print(
                    f"Failed to generate resume at index {index}")
                print("" + "-" * 50)
                continue

            print(
                f"Generated resume for at index {index} Resume link: {resume_link}")

            job["file_url"] = resume_link

            all_jobs[index] = job

            with open(path, "w", encoding="utf-8") as f:
                json.dump(all_jobs, f, indent=4)

            print("" + "-" * 50)
        else:
            print(
                f"Already generated resume for at index {index}")
            print("" + "-" * 50)
            continue

    print("" + "-" * 50)

    print(f"Updated jobs.json with resume links.")


if __name__ == "__main__":
    main()
